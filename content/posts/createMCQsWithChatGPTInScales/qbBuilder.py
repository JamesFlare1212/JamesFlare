import argparse
import json
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define command-line arguments
parser = argparse.ArgumentParser(description='Generate test papers, answer sheets, and marking scheme.')
parser.add_argument('-q', '--question-bank', type=str, required=True, help='Path to JSON file containing question bank.')
parser.add_argument('-n', '--number-question', type=int, required=True, help='Number of questions to include in the test paper.')
args = parser.parse_args()

# Open JSON file and load data
with open(args.question_bank, 'r') as f:
    data = json.load(f)

# Select random questions
selected_questions = random.sample(data, args.number_question)

# Create Test Paper, Answer Sheet, and Marking Scheme
test_paper_doc = Document()
answer_sheet_doc = Document()
marking_scheme_doc = Document()

# Set font
font = test_paper_doc.styles['Normal'].font
font.name = 'Arial'
font.size = Pt(12)

# Add test paper title
test_paper_title = test_paper_doc.add_paragraph('Test Paper', style='Title')
test_paper_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
test_paper_doc.add_paragraph()

# Add answer sheet title
answer_sheet_title = answer_sheet_doc.add_paragraph('Answer Sheet', style='Title')
answer_sheet_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
answer_sheet_doc.add_paragraph()

# Add marking scheme title
marking_scheme_title = marking_scheme_doc.add_paragraph('Marking Scheme', style='Title')
marking_scheme_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
marking_scheme_doc.add_paragraph()

# Loop through selected questions and add to documents
for i, question in enumerate(selected_questions):
    # Test Paper
    test_paper_doc.add_paragraph(f'{i+1}. {question["question"]}')
    test_paper_doc.add_paragraph(f'     A. {question["option_a"]}')
    test_paper_doc.add_paragraph(f'     B. {question["option_b"]}')
    test_paper_doc.add_paragraph(f'     C. {question["option_c"]}')
    test_paper_doc.add_paragraph(f'     D. {question["option_d"]}')
    test_paper_doc.add_paragraph()

    # Answer Sheet
    answer_sheet_doc.add_paragraph(f'{i+1}. {question["question"]}')
    answer_sheet_doc.add_paragraph(f'     A. {question["option_a"]}')
    answer_sheet_doc.add_paragraph(f'     B. {question["option_b"]}')
    answer_sheet_doc.add_paragraph(f'     C. {question["option_c"]}')
    answer_sheet_doc.add_paragraph(f'     D. {question["option_d"]}')
    answer_sheet_doc.add_paragraph(f'Answer: {question["answer"]}')
    answer_sheet_doc.add_paragraph(f'Explanation: {question["explanation"]}')
    answer_sheet_doc.add_paragraph()

    # Marking Scheme
    marking_scheme_doc.add_paragraph(f'{i+1}. {question["question"]}')
    marking_scheme_doc.add_paragraph(f'Answer: {question["answer"]}')
    marking_scheme_doc.add_paragraph()

# Save documents
test_paper_doc.save('test_paper.docx')
answer_sheet_doc.save('answer_sheet.docx')
marking_scheme_doc.save('marking_scheme.docx')